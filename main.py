import os
import time
import asyncio
import io
import uuid
from typing import Optional
from fastapi import FastAPI, Form, File, UploadFile, HTTPException, Header, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel
from dotenv import load_dotenv

from copyleaks.copyleaks import Copyleaks
from copyleaks.models.submit.ai_detection_document import NaturalLanguageDocument

# ------------------------------------------------------------
# Load environment variables
# ------------------------------------------------------------
load_dotenv()

EMAIL_ADDRESS = os.getenv("COPyleAKS_EMAIL")
API_KEY = os.getenv("COPyleAKS_API_KEY")
ALLOWED_ORIGIN = os.getenv("ALLOWED_ORIGIN", "http://localhost:3000")
MAX_FILE_SIZE = int(os.getenv("MAX_FILE_SIZE", 5 * 1024 * 1024))  # 5 MB
AUTH_TOKEN = os.getenv("API_AUTH_TOKEN")  # optional; if set, clients must send it

if not EMAIL_ADDRESS or not API_KEY:
    raise ValueError("Missing Copyleaks credentials in environment variables")

# ------------------------------------------------------------
# FastAPI app & static serving
# ------------------------------------------------------------
app = FastAPI(title="Team-5 AI Detection API")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
STATIC_DIR = os.path.join(BASE_DIR, "static")
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

@app.get("/")
async def root():
    index_path = os.path.join(STATIC_DIR, "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path)
    return {"message": "Welcome to Team-5 API"}

# ------------------------------------------------------------
# CORS – restricted to a single origin
# ------------------------------------------------------------
app.add_middleware(
    CORSMiddleware,
    allow_origins=[ALLOWED_ORIGIN],
    allow_methods=["GET", "POST"],
    allow_headers=["*"],
)

# ------------------------------------------------------------
# Token cache for Copyleaks
# ------------------------------------------------------------
_token_cache = {"token": None, "expires_at": 0}

def get_auth_token():
    """Authenticate with Copyleaks and cache the token."""
    now = time.time()
    # Refresh if token is missing or will expire within 60 seconds
    if _token_cache["token"] and _token_cache["expires_at"] > now + 60:
        return _token_cache["token"]

    try:
        token = Copyleaks.login(EMAIL_ADDRESS, API_KEY)
        # Assume token lifetime is 1 hour; cache for 50 minutes
        _token_cache["token"] = token
        _token_cache["expires_at"] = now + 50 * 60
        return token
    except Exception as e:
        # Log the error internally, but return a generic message
        print(f"Copyleaks login failed: {e}")
        raise HTTPException(status_code=503, detail="Copyleaks authentication service unavailable")

# ------------------------------------------------------------
# Optional API authentication (via Bearer token)
# ------------------------------------------------------------
def verify_auth(authorization: Optional[str] = Header(None)):
    if AUTH_TOKEN:
        if not authorization or authorization != f"Bearer {AUTH_TOKEN}":
            raise HTTPException(status_code=401, detail="Unauthorized")
    return True

# ------------------------------------------------------------
# Pydantic response model
# ------------------------------------------------------------
class DetectionResponse(BaseModel):
    percentage: int
    ai_percent: int
    human_percent: int
    model_version: Optional[str] = None

# ------------------------------------------------------------
# Text extraction helpers (unchanged, but with added error safety)
# ------------------------------------------------------------
def extract_text_from_txt(raw_bytes: bytes) -> str:
    try:
        return raw_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return raw_bytes.decode("latin-1")

def extract_text_from_docx(raw_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(raw_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)

def extract_text_from_pdf(raw_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(raw_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)

async def extract_text_from_file(file: UploadFile) -> str:
    """
    Extract text from uploaded file.
    Checks file size, extension, and (when available) MIME type.
    """
    # Read content
    raw_bytes = await file.read()
    if len(raw_bytes) == 0:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")
    if len(raw_bytes) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Maximum allowed size is {MAX_FILE_SIZE // (1024*1024)} MB"
        )

    filename = (file.filename or "").lower()
    # Optional MIME type check (if provided by client)
    content_type = file.content_type or ""
    allowed_mimes = {
        "text/plain": ".txt",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "application/pdf": ".pdf",
    }
    if content_type and content_type not in allowed_mimes:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported MIME type: {content_type}. Allowed: txt, docx, pdf"
        )

    # Fallback to extension check
    try:
        if filename.endswith(".txt") or content_type == "text/plain":
            return extract_text_from_txt(raw_bytes)
        elif filename.endswith(".docx") or content_type == allowed_mimes[".docx"]:
            return extract_text_from_docx(raw_bytes)
        elif filename.endswith(".pdf") or content_type == "application/pdf":
            return extract_text_from_pdf(raw_bytes)
        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type. Please upload .txt, .docx, or .pdf"
            )
    except HTTPException:
        raise
    except Exception as e:
        # Log the error but return a generic message
        print(f"File extraction error: {e}")
        raise HTTPException(status_code=400, detail="Could not read file content")

# ------------------------------------------------------------
# Copyleaks detection runner (sanitised error handling)
# ------------------------------------------------------------
def run_detection(text: str, scan_id: str):
    """Run AI detection using Copyleaks SDK. Synchronous – called via run_in_executor."""
    auth_token = get_auth_token()
    document = NaturalLanguageDocument(text)
    # document.set_sandbox(True)   # uncomment if using sandbox
    try:
        return Copyleaks.AiDetectionClient.submit_natural_language(
            auth_token, scan_id, document
        )
    except Exception as e:
        # Log full details for debugging
        print("=" * 60)
        print(f"Copyleaks detection error: {e}")
        import traceback
        traceback.print_exc()
        print("=" * 60)
        # Raise a generic HTTP exception – internal details are not leaked
        raise HTTPException(
            status_code=502,
            detail="AI detection service returned an error. Please try again later."
        )

# ------------------------------------------------------------
# Main detection endpoint
# ------------------------------------------------------------
@app.post("/detect", response_model=DetectionResponse)
async def detect_ai(
    auth_ok: bool = Depends(verify_auth),   # optional auth
    text: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
):
    # --- 1. Reject if both text and file are provided ---
    if file is not None and file.filename and text is not None:
        raise HTTPException(status_code=400, detail="Provide either text or a file, not both.")

    # --- 2. Resolve input ---
    if file is not None and file.filename:
        resolved_text = await extract_text_from_file(file)
    elif text is not None:
        resolved_text = text
    else:
        raise HTTPException(status_code=400, detail="No text or file provided")

    # --- 3. Validate content ---
    stripped = resolved_text.strip()
    if not stripped:
        raise HTTPException(status_code=400, detail="No text content found to analyze")
    if len(stripped) < 350:
        raise HTTPException(
            status_code=400,
            detail="Text must be at least 350 characters (recommend 500+ for accuracy)."
        )

    # --- 4. Run detection ---
    scan_id = str(uuid.uuid4())
    try:
        loop = asyncio.get_event_loop()
        response = await loop.run_in_executor(None, run_detection, resolved_text, scan_id)

        # Log raw response internally (for debugging) – but be cautious about sensitive data
        # print("Copyleaks response:", json.dumps(response, indent=2))  # optional

        # Safely extract fields
        summary = response.get("summary", {})
        ai_score = int(summary.get("ai", 0) * 100)
        human_score = int(summary.get("human", 0) * 100)
        model_version = response.get("modelVersion")

        # Business logic: scores below 20% are rounded down to 0
        final_score = 0 if ai_score < 20 else ai_score

        return DetectionResponse(
            percentage=final_score,
            ai_percent=ai_score,
            human_percent=human_score,
            model_version=model_version,
        )

    except HTTPException:
        raise
    except Exception as e:
        # Catch-all for any unexpected errors – log and return generic
        print(f"Unexpected error in /detect: {e}")
        raise HTTPException(status_code=500, detail="Internal server error. Please try again later.")
